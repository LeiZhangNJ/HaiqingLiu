#!/usr/bin/env python
# coding: utf-8

# In[2]:


import os
import numpy as np
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.naive_bayes import MultinomialNB
from sklearn import metrics
from sklearn.naive_bayes import GaussianNB
from sklearn.naive_bayes import BernoulliNB


# In[3]:


def loadfile(filepath,leibie):
    
    filelist = os.listdir(filepath)
    content = []
    label = []
    
    for file in filelist:
        with open(filepath+"/"+file,encoding='utf-8')as f:
            content.append("".join((f.read())))
            label.append(leibie)
    return content,label


# In[4]:


wenbendir = ['train','test']
currentdir = os.getcwd()
traincontent = []
trainlabel = []
testcontent = []
testlabel = []
for wenben in wenbendir:#['train','test']
#     os.chdir(wenben)
    wenbenlist = os.listdir(wenben)#【相关 不相关】
    for leibie in wenbenlist:
        content,label=loadfile(wenben+"/"+leibie,leibie)
        if wenben=="train":
            traincontent += content
            trainlabel +=label
        elif wenben=='test':
            testcontent += content
            testlabel +=label
        os.chdir(currentdir)


# In[5]:


with open(r'F:\zhuomian\数据库\stopword.txt',encoding='utf-8')as file:
    stopwords = file.read().split("\n")


# In[6]:


tfidf = TfidfVectorizer(stop_words=stopwords,max_df=0.5)
traindata = tfidf.fit_transform(traincontent)
testdata = tfidf.transform(testcontent)


# In[7]:


from sklearn.linear_model import LinearRegression,LogisticRegression,Ridge,RidgeCV,Lasso, LassoCV
from sklearn.model_selection import train_test_split,GridSearchCV,cross_val_score,cross_validate
from sklearn import  metrics as mt
from  statsmodels.stats.outliers_influence import variance_inflation_factor
from sklearn.metrics import classification_report
# 逻辑回归 LogisticRegression
model = LogisticRegression()
model.fit(traindata,trainlabel)
predict_test = model.predict(testdata)
print(predict_test)
print("逻辑回归文本分类的准确率为：",metrics.accuracy_score(predict_test,testlabel))


# In[14]:


# 混淆矩阵
from sklearn.metrics import confusion_matrix
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sn
cm = confusion_matrix(testlabel,predict_test)
df_cm = pd.DataFrame(cm)
df_cm
#annot = True 显示数字 ，fmt参数不使用科学计数法进行显示
ax = sn.heatmap(df_cm,annot=True,fmt='.20g')
ax.set_title('confusion matrix') #标题
ax.set_xlabel('predict') #x轴
ax.set_ylabel('true') #y轴


# ax.savefig('混淆矩阵.jpg')
# plt.figure(figsize=(4, 4), dpi=144)
# plt.title('Confusion matrix of the classifier')
# ax = plt.gca()                                  
# ax.spines['right'].set_color('none')            
# ax.spines['top'].set_color('none')
# ax.spines['bottom'].set_color('none')
# ax.spines['left'].set_color('none')
# ax.xaxis.set_ticks_position('none')
# ax.yaxis.set_ticks_position('none')
# ax.set_xticklabels([])
# ax.set_yticklabels([])
# plt.matshow(cm, fignum=1, cmap='gray')
# plt.colorbar();


# In[9]:


#多项式朴素贝叶斯
nb_model = MultinomialNB(alpha=0.001)
nb_model.fit(traindata,trainlabel)
predict_test = nb_model.predict(testdata)
print("多项式朴素贝叶斯文本分类的准确率为：",metrics.accuracy_score(predict_test,testlabel))


# In[27]:


#网格搜索交叉验证
from sklearn.model_selection import GridSearchCV

# alpha_list= np.linspace(0.01,0.001,100)
alpha_list=[0.01,0.001,0.0001]

# Set the parameters by cross-validation
param_grid = [{'alpha': alpha_list, }]

clf = GridSearchCV(MultinomialNB(), param_grid, cv=5)
clf.fit(traindata,trainlabel)
print("best param: {0}\nbest score: {1}".format(clf.best_params_, 
                                                clf.best_score_))


# In[28]:


#bernoulli朴素贝叶斯
from sklearn.naive_bayes import BernoulliNB
ber_model = BernoulliNB(alpha=0.001)
ber_model.fit(traindata,trainlabel)
ber_predict = ber_model.predict(testdata)
print("bernoulli贝叶斯文本分类的准确率为：",metrics.accuracy_score(ber_predict,testlabel))


# In[29]:


#高斯贝叶斯分类器
gauss_model = GaussianNB()
gauss_model.fit(traindata.toarray(),trainlabel)
gauss_predict = ber_model.predict(testdata.toarray())
print("GaussianNB贝叶斯文本分类的准确率为：",metrics.accuracy_score(gauss_predict,testlabel))


# In[1]:


# 创建文件
import os
 
def mkdir(path):
	folder = os.path.exists(path)
	if not folder:                   #判断是否存在文件夹如果不存在则创建为文件夹
		os.makedirs(path)            #makedirs 创建文件时如果路径不存在会创建这个路径
		print("---  new folder...  ---")
		print("---  OK  ---")
 
	else:
		print("---  There is this folder!  ---")
		

years = 2020
while(1):
    mkpath1=r"F:\zhuomian\新建文件夹 (10)\{}".format(str(years))
    mkpath2=r"F:\zhuomian\新建文件夹 (10)\{}筛选".format(str(years))
# 调用函数
    mkdir(mkpath1)
    mkdir(mkpath2)
    print(years)
    years= years - 1
    if years<1900:
        break


# In[10]:


# 筛选文件
import docx
import re
import os 
import shutil


years = 2020
while(1):
        file=docx.Document(r"F:\zhuomian\新建文件夹 (10)\{years}.docx".format(years = years))
        a = 1
        for para in file.paragraphs:
                path=r'F:\zhuomian\新建文件夹 (10)\{years}\{number}.txt'.format(years = years,number = str(a))
                f = open(path, 'a+', encoding='UTF8')
                print(para.text,file=f)
                a = a + 1
        #d
        file_path = 'F:\zhuomian\新建文件夹 (10)\{years}'.format(years = str(years))
        filelist = os.listdir(file_path)
        content = []
        for file in filelist:
                with open(file_path+"/"+file,encoding='utf-8')as f:
                    content.append("".join((f.read())))
        data = tfidf.transform(content)
        predict = model.predict(data)
        #d
        number = 1
        for i in predict:

            if i == '相关':
                full_path = file_path+"/"+'{}.txt'.format(number)
                des_path ='F:\zhuomian\新建文件夹 (10)\{years}筛选'.format(years = str(years))
                shutil.move(full_path, des_path)

        #         print(number)
                number =number+1
            else:
                number =number+1
                continue

        years=years - 1
        if years<=1900:
                    break


# In[12]:


# 把筛选出来的文件整合到一个文件
import os,shutil
years = 1962
while(1):
    try:
        path = r"F:\zhuomian\新建文件夹 (10)\{}筛选".format(years)
        f=open(r'F:\zhuomian\新建文件夹 (11)\{}.txt'.format(years),"a",encoding='UTF-8')
        for r,d,fi in os.walk(path):
            for files in fi:
                     if files.endswith(".txt"):
                          path2=os.path.join(r,files)
                          g=open(r'{}'.format(path2),encoding='UTF-8')
                          shutil.copyfileobj(g,f)
                          g.close()
        f.close()
        print(years)
        years= years - 1
        if years<1900:
            break
    except:
        break


# In[48]:


path = os.path.join("F:\\zhuomian\\新建文件夹 (7)\\1941",'1.txt')
r'{}'.format(path)


# In[18]:


years = '2020'


# In[16]:


import docx
import re

# years = '一九九五'
#获取文档
file=docx.Document(r"F:\zhuomian\实验数据TRYONE\{years}.docx".format(years = years))
print("段落数:"+str(len(file.paragraphs))) #输出段落数
# file_word = docx.Document()
a = 1
#输出每一段的内容
for para in file.paragraphs:
    path=r'F:\zhuomian\实验数据TRYONE\{years}\{number}.txt'.format(years = years,number = str(a))
    f = open(path, 'a+', encoding='UTF8')
    print(para.text,file=f)
    a = a + 1


# In[8]:


# 预测存储
import os 
import shutil
file_path = 'F:\zhuomian\实验数据TRYONE\{years}'.format(years = years)
filelist = os.listdir(file_path)
content = []
for file in filelist:
        with open(file_path+"/"+file,encoding='utf-8')as f:
            content.append("".join((f.read())))
data = tfidf.transform(content)
predict = model.predict(data)
# print(predict)
# for i in predict:
#     number = 1
#     if i == '相关':
#         full_path = file_path+"/"+'{}.txt'.format(number)
#         des_path ='F:\zhuomian\新建文件夹 (2)'
#         shutil.move(full_path, des_path)
#     else:
#         continue
        
#     number =number+1



# In[9]:


number = 1
for i in predict:
    
    if i == '相关':
        full_path = file_path+"/"+'{}.txt'.format(number)
        des_path ='F:\zhuomian\实验数据TRYONE\{years}筛选'.format(years = years)
        shutil.move(full_path, des_path)
        
#         print(number)
        number =number+1
    else:
        number =number+1
        continue
        
    


# In[ ]:




